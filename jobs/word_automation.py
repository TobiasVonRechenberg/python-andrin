""" Source: https://python-docx.readthedocs.io/en/latest/ """

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Poststrasse 18 111 Niederhausen', 0)

p = document.add_paragraph('28 November, 2019')
('\n')
p.add_run('                                                                                                                                                                                                                                                                ')
p.add_run(' von Herrn ')
p.add_run('Müller.').Müller = True

document.add_heading(' An Frau Müller, 99', level=1)
document.add_paragraph('Betrag der Rechnung', style='Intense Quote')

document.add_paragraph(
    'Erste Spalte = Kosten', style='List Number'
)
document.add_paragraph(
    'Zweite Liste = Versandkosten', style='List Number'
)

document.add_picture('../input/Bildgeld.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('../output/demo.docx')
